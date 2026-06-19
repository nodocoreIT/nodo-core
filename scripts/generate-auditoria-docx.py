#!/usr/bin/env python3
"""Generate auditoria.docx for nodo-inmo marketing vs implementation audit."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "auditoria.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build() -> None:
    doc = Document()

    title = doc.add_heading("Auditoría nodo-inmo", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Comparación entre lo prometido en la web de precios y lo implementado en el código"
    )
    r.italic = True
    r.font.size = Pt(11)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(
        "Fuente web: https://www.nodocore.com.ar/nodo-inmo#precios\n"
        "Repositorio: nodo-core / nodo-inmo\n"
        "Fecha: junio 2026"
    )
    m.font.size = Pt(10)
    m.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    add_heading(doc, "Veredicto general", 1)
    add_table(
        doc,
        ["Plan", "Prometido en web", "Entregado hoy", "Gap principal"],
        [
            [
                "Starter",
                "13 ítems",
                "~7 sólidos, 6 parciales",
                "Varios ítems incompletos o mal ubicados en el plan",
            ],
            [
                "Pro",
                "8 extras",
                "~2–3 con código real",
                "Integraciones, bot 24/7, stats y NODO ID ausentes",
            ],
        ],
    )
    add_para(
        doc,
        "Riesgo comercial: un cliente Starter puede usar bien propiedades, contratos básicos, "
        "cobros y caja. Un cliente Pro que espere integraciones, bot 24/7, estadísticas por "
        "empleado o NODO ID no los va a encontrar en el producto actual.",
    )

    add_heading(doc, "Plan Starter — ítem por ítem", 1)
    add_table(
        doc,
        ["Feature (web)", "Estado", "Qué hace realmente"],
        [
            [
                "Alta y ficha completa de propiedad",
                "✅ Implementado",
                "CRUD en /admin/properties, ficha rica (dirección, tipo, precio, m², amenities, fotos). "
                "Detalle en modal, no URL propia /properties/:id.",
            ],
            [
                "Fotos y documentos adjuntos",
                "⚠️ Parcial",
                "Fotos: sí (storage + galería). Documentos: tabla y upload en /admin/documentos, "
                "pero no integrados en la ficha de la propiedad.",
            ],
            [
                "Estados disponible / alquilada / vendida",
                "✅ Implementado",
                "available, rented, sold (+ reserved, negotiation, inactive).",
            ],
            [
                "Búsqueda y filtros avanzados",
                "✅ Implementado",
                "Filtros por operación, tipo, precio, provincia, localidad, propietario, ambientes + búsqueda global.",
            ],
            [
                "Web interna con detalle de propiedad",
                "⚠️ Parcial / mal plan",
                "Existe PortalPage (catálogo interno), pero está bloqueado a Pro (PlanGate). "
                "En Starter solo hay el editor modal.",
            ],
            [
                "Cálculo automático aumentos ICL/IPC",
                "⚠️ Parcial",
                "Se configura índice y fecha en el contrato; dashboard próximos aumentos; badge IPC. "
                "No recalcula el alquiler ni regenera cuotas automáticamente.",
            ],
            [
                "Alertas vencimiento de contrato",
                "⚠️ Parcial",
                "Alertas en el dashboard admin (configurables). Sin email/WhatsApp automático.",
            ],
            [
                "Cobros efectivo y transferencia",
                "⚠️ Parcial",
                "El schema admite cash y transfer, pero el diálogo de cobro siempre guarda transfer (hardcodeado). "
                "Efectivo no se puede elegir en UI.",
            ],
            [
                "Cuentas bancarias y caja chica",
                "✅ Implementado",
                "Cuentas en settings, movimientos en /admin/caja, conceptos, saldos. Solo admin (no agentes).",
            ],
            [
                "Historial de pagos e informe de morosidad",
                "⚠️ Parcial",
                "Historial de cuotas en /admin/payments con filtros overdue/pending/paid. "
                "No hay informe de morosidad (PDF/export dedicado).",
            ],
            [
                "Pipeline ventas (interesado, reserva)",
                "⚠️ Parcial",
                "Estados reserved / negotiation en propiedades. No hay pipeline (etapas, leads, tablero).",
            ],
            [
                "Roles Admin y Agentes",
                "✅ Implementado",
                "Invitación, JWT, routing. Agentes sin caja/rendiciones/ganancias.",
            ],
            [
                "Acceso web y móvil, sin instalación",
                "✅ Implementado",
                "SPA responsive (Vite + React), sin app nativa.",
            ],
        ],
    )

    add_heading(doc, "Starter — lo que funciona bien", 2)
    for item in [
        "Gestión de propiedades de punta a punta",
        "Contratos de alquiler (alta, cuotas, cobro)",
        "Caja y cuentas bancarias",
        "Roles admin/agente",
        "Filtros y búsqueda",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Starter — no cumple lo vendido", 2)
    gaps_starter = [
        "“Cálculo automático ICL/IPC” → solo configuración + avisos, no cálculo",
        "“Informe de morosidad” → listados overdue, no informe",
        "“Pipeline interesado/reserva” → solo estados en la propiedad",
        "“Cobros efectivo” → no selectable en UI",
        "“Web interna con detalle” → implementada pero solo Pro",
    ]
    for g in gaps_starter:
        doc.add_paragraph(g, style="List Bullet")

    add_heading(doc, "Plan Pro — ítem por ítem", 1)
    add_table(
        doc,
        ["Feature (web)", "Estado", "Qué hace realmente"],
        [
            [
                "Portal Propietario",
                "⚠️ Parcial",
                "Rutas /owner/*, propiedades y liquidaciones. RLS por contacts.portal_user_id. "
                "La app no vincula portal_user_id al invitar → portal puede quedar vacío sin seed manual.",
            ],
            [
                "Portal Inquilinos (contrato, pagos, reclamos)",
                "⚠️ Parcial",
                "/tenant/contrato, /tenant/pagos (solo lectura), /tenant/reclamos (end-to-end). Sin pago online.",
            ],
            [
                "Avisos vencimiento, aumentos y mora",
                "⚠️ Parcial",
                "Dashboard + campana de notificaciones admin. WhatsApp manual para aumentos. "
                "Bot automático marcado coming_soon en código.",
            ],
            [
                "Estadísticas ventas por empleado",
                "❌ No existe",
                "Solo tarjeta coming_soon en automation-definitions.ts.",
            ],
            [
                "Gmail, Google Sheets, Mercado Pago",
                "❌ No existe",
                "Los tres en coming_soon. Página de automatizaciones sin ruta en el router admin.",
            ],
            [
                "Administración automática redes sociales",
                "❌ No existe",
                "social-auto → coming_soon.",
            ],
            [
                "Generación automática de contratos",
                "⚠️ Parcial",
                "PDF de locación sí (generador real). No es automático al cargar propiedad+partes; "
                "requiere contrato existente + click.",
            ],
            [
                "NODO ID",
                "❌ No existe en app",
                "Columna shared.nodo_id en DB (placeholder). Cero UI ni uso en frontend.",
            ],
        ],
    )

    add_heading(doc, "Pro — lo más sólido hoy", 2)
    for item in [
        "Reclamos inquilino ↔ admin",
        "PDF de contrato de locación",
        "Portales owner/tenant (estructura + RLS), si el contacto está vinculado",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Pro — bloqueadores", 2)
    for item in [
        "Vincular portal_user_id al invitar propietario/inquilino",
        "Enforcement de plan (JWT no trae tier; PlanGate solo en /admin/portal)",
        "Integraciones y automatizaciones — marketing puro por ahora",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Automatizaciones en código (coming_soon)", 1)
    add_table(
        doc,
        ["ID", "Título", "Estado"],
        [
            ["wa-bot-queries", "Respuesta automática por WhatsApp", "coming_soon"],
            ["wa-expiry-alerts", "Avisos de vencimiento y mora", "coming_soon"],
            ["auto-contract", "Generación automática de contrato", "active (parcial)"],
            ["employee-stats", "Estadísticas de ventas por empleado", "coming_soon"],
            ["mercadopago", "Cobros online con Mercado Pago", "coming_soon"],
            ["gmail-integration", "Integración con Gmail", "coming_soon"],
            ["google-sheets", "Sincronización con Google Sheets", "coming_soon"],
            ["social-auto", "Publicación automática en redes sociales", "coming_soon"],
        ],
    )
    add_para(
        doc,
        "Archivo fuente: nodo-inmo/src/features/automations/automation-definitions.ts",
        bold=False,
    )

    add_heading(doc, "Desalineaciones graves web ↔ producto", 1)
    add_table(
        doc,
        ["En la web decís", "En el código pasa"],
        [
            ["Starter: web interna con detalle", "Catálogo interno = Pro"],
            ["Starter: cálculo automático ICL/IPC", "Solo config + alertas"],
            ["Starter: informe de morosidad", "Filtro vencido, sin informe"],
            ["Starter: cobros efectivo", "Solo transferencia en UI"],
            ["Pro: bot 24/7 + automatizaciones", "coming_soon"],
            ["Pro: Gmail / Sheets / MP", "coming_soon"],
            ["Pro: redes automáticas", "coming_soon"],
            ["Pro: stats por empleado", "coming_soon"],
            ["Pro: NODO ID", "Tabla vacía, sin UI"],
        ],
    )

    add_heading(doc, "Matriz resumida", 1)
    add_para(doc, "STARTER", bold=True)
    for line in [
        "✅ Propiedades, estados, filtros, roles, web responsive",
        "✅ Contratos, cobros (transfer), caja, historial pagos",
        "⚠️ Fotos+docs, alertas, ICL/IPC, mora, pipeline, web interna",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_para(doc, "PRO (además de Starter)", bold=True)
    for line in [
        "✅ Reclamos, PDF contrato, portales (con caveats)",
        "⚠️ Portales onboarding, avisos automáticos, auto contrato",
        "❌ MP, Gmail, Sheets, redes, stats empleado, NODO ID",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "Prioridades recomendadas", 1)
    add_heading(doc, "Corto plazo (honestidad + Starter usable)", 2)
    for item in [
        "Corregir copy o mover web interna al plan correcto",
        "Selector efectivo/transferencia en cobros",
        "Informe morosidad básico (export/PDF de overdue)",
        "Aclarar en web que ICL/IPC es seguimiento, no recálculo automático (o implementarlo)",
    ]:
        doc.add_paragraph(item, style="List Number")

    add_heading(doc, "Medio plazo (Pro vendible)", 2)
    for item in [
        "portal_user_id al invitar propietario/inquilino",
        "Plan gating real en JWT + rutas",
        "Un canal de avisos automáticos (WhatsApp o email)",
    ]:
        doc.add_paragraph(item, style="List Number")

    add_heading(doc, "Largo plazo (ecosistema)", 2)
    for item in [
        "Mercado Pago",
        "NODO ID",
        "Integraciones (Gmail, Sheets)",
        "Estadísticas por empleado",
    ]:
        doc.add_paragraph(item, style="List Number")

    add_heading(doc, "Evidencia técnica (rutas y archivos clave)", 1)
    add_table(
        doc,
        ["Área", "Rutas / archivos"],
        [
            ["Propiedades", "src/features/properties/, /admin/properties"],
            ["Fotos", "supabase/migrations/20260606140000_*, photo-gallery.tsx"],
            ["Documentos", "src/features/documentos/, 20260606130000_create_documents.sql"],
            ["Contratos / ICL-IPC", "src/features/contracts/, src/features/ipc/"],
            ["Cobros", "src/features/payments/payment-collect-dialog.tsx (transfer hardcoded)"],
            ["Caja", "src/features/caja/, /admin/caja"],
            ["Portal interno (Pro)", "src/features/portal/, PlanGate en admin-portal-page.tsx"],
            ["Portal propietario", "src/portals/owner/"],
            ["Portal inquilino", "src/portals/tenant/"],
            ["Reclamos", "src/features/reclamos/"],
            ["Automatizaciones", "src/features/automations/automation-definitions.ts"],
            ["PDF contrato", "src/features/contracts/ (contract-locacion-*)"],
            ["NODO ID", "shared.nodo_id en migrations, sin UI"],
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Documento generado para Nodo Core — uso interno")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(OUTPUT)
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    build()
